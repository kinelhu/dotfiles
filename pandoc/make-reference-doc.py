#!/usr/bin/env python3
"""
Generate a pandoc reference.docx matching the LaTeX/PDF style (stats-report-template.tex).

Colors  : AccentPrimary #27A9BC, TextPrimary #1A1A1A, TextSecondary #949494
Font    : Inter (body/headings), JetBrains Mono (code)
Margins : top/bottom 2.5 cm, left/right 3 cm  (US Letter, portrait)
"""

import os, subprocess, tempfile
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DOCX   = "/tmp/base-reference.docx"
OUTPUT_PATH = os.path.expanduser("~/.local/share/pandoc/reference.docx")

# ── Palette ──────────────────────────────────────────────────────────────────
ACCENT   = RGBColor(0x27, 0xA9, 0xBC)   # #27A9BC  teal
PRIMARY  = RGBColor(0x1A, 0x1A, 0x1A)   # #1A1A1A  quasi-noir
SECONDARY= RGBColor(0x94, 0x94, 0x94)   # #949494  gris moyen
BORDER   = RGBColor(0xDD, 0xDD, 0xDD)   # #DDDDDD  séparateurs légers

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(style, name="Arial", size_pt=10.5, bold=False,
             italic=False, color=None):
    """Set font properties on a style, covering all Unicode ranges."""
    font = style.font
    font.size  = Pt(size_pt)
    font.bold  = bold
    font.italic = italic
    if color:
        font.color.rgb = color

    # Force the font for ASCII + hAnsi + complex-script (covers accented French)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)


def set_para(style, before=0, after=6, line_spacing=1.15, keep_next=False):
    pf = style.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing      = line_spacing
    pf.first_line_indent = Pt(0)
    pf.left_indent       = Pt(0)
    if keep_next:
        pf.keep_with_next = True


def remove_numbering(style):
    """Strip any automatic numbering from a heading style."""
    pPr = style.element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def add_footer(section):
    """Three-part footer: today left | page centre | author right."""
    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    p = footer.paragraphs[0]
    p.clear()

    # Tab stops: centre and right
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    for align, pos in (("center", 4677), ("right", 9354)):
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"),   align)
        tab.set(qn("w:pos"),   str(pos))
        tabs.append(tab)
    pPr.append(tabs)

    def run(text=None, field=None):
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz");    sz.set(qn("w:val"), "18")   # 9 pt
        szCs= OxmlElement("w:szCs"); szCs.set(qn("w:val"), "18")
        col = OxmlElement("w:color"); col.set(qn("w:val"), "949494")
        rpr.extend([sz, szCs, col])
        r.append(rpr)
        if field:
            fldBegin = OxmlElement("w:fldChar")
            fldBegin.set(qn("w:fldCharType"), "begin")
            inst     = OxmlElement("w:instrText")
            inst.text = field
            fldSep   = OxmlElement("w:fldChar")
            fldSep.set(qn("w:fldCharType"), "separate")
            fldEnd   = OxmlElement("w:fldChar")
            fldEnd.set(qn("w:fldCharType"), "end")
            r.append(fldBegin)
            r2 = OxmlElement("w:r"); r2.append(fldSep)
            r3 = OxmlElement("w:r"); r3.append(fldEnd)
            p._p.append(r)
            p._p.append(r2)
            p._p.append(r3)
            return
        t = OxmlElement("w:t"); t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p._p.append(r)

    def tab():
        r = OxmlElement("w:r")
        t = OxmlElement("w:tab")
        r.append(t)
        p._p.append(r)

    run(field=" DATE \\@ \"d MMM yyyy\" \\* MERGEFORMAT")
    tab()
    run(field=" PAGE \\* MERGEFORMAT")
    tab()
    run("K. El Husseini")

# ── Main ──────────────────────────────────────────────────────────────────────

doc = Document(BASE_DOCX)

# Page geometry
sec = doc.sections[0]
sec.page_width   = Inches(8.5)
sec.page_height  = Inches(11)
sec.top_margin   = Cm(2.5)
sec.bottom_margin= Cm(2.5)
sec.left_margin  = Cm(3.0)
sec.right_margin = Cm(3.0)

s = doc.styles

# Normal (body text)
set_font(s["Normal"], size_pt=10.5)
set_para(s["Normal"], before=0, after=6, line_spacing=1.15)

# Heading 1 — \Large\bfseries\color{AccentPrimary}
set_font(s["Heading 1"], size_pt=16, bold=True, color=ACCENT)
set_para(s["Heading 1"], before=14, after=6, line_spacing=1.0, keep_next=True)
remove_numbering(s["Heading 1"])

# Heading 2 — \large\bfseries\color{TextPrimary}
set_font(s["Heading 2"], size_pt=13, bold=True, color=PRIMARY)
set_para(s["Heading 2"], before=10, after=5, line_spacing=1.0, keep_next=True)
remove_numbering(s["Heading 2"])

# Heading 3 — \normalsize\bfseries\color{TextSecondary}
set_font(s["Heading 3"], size_pt=11, bold=True, color=SECONDARY)
set_para(s["Heading 3"], before=8, after=4, line_spacing=1.0, keep_next=True)
remove_numbering(s["Heading 3"])

# Heading 4-6 (inherited styling, just clean numbering)
for lvl in ("Heading 4", "Heading 5", "Heading 6"):
    if lvl in [st.name for st in s]:
        set_font(s[lvl], size_pt=10.5, bold=True)
        remove_numbering(s[lvl])

# Body Text / Compact
for name in ("Body Text", "First Paragraph"):
    if name in [st.name for st in s]:
        set_font(s[name], size_pt=10.5)
        set_para(s[name], before=0, after=6, line_spacing=1.15)

if "Compact" in [st.name for st in s]:
    set_font(s["Compact"], size_pt=10.5)
    set_para(s["Compact"], before=0, after=2, line_spacing=1.0)

# Verbatim / code
for name in ("Verbatim Char", "Source Code", "Verbatim"):
    if name in [st.name for st in s]:
        set_font(s[name], name="Courier New", size_pt=9.0)

# Footer
add_footer(sec)

doc.save(OUTPUT_PATH)
print(f"✓  Saved → {OUTPUT_PATH}")
